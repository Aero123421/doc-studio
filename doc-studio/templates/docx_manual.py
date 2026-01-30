#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
マニュアル/ドキュメント Wordテンプレート
- 技術文書スタイル
- カラー: ブルー系
"""

import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _load_data(args: argparse.Namespace) -> dict:
    if args.data and args.data_file:
        raise SystemExit("Use either --data or --data-file (not both)")

    if args.data_file:
        p = Path(args.data_file)
        return json.loads(p.read_text(encoding="utf-8"))

    if args.data:
        return json.loads(args.data)

    return {}

def set_cell_shading(cell, color):
    """セルの背景色を設定"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, code_text):
    """コードブロックを追加"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 背景色をグレーに
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'f5f5f5')
    p._p.get_or_add_pPr().append(shading_elm)

def create_manual(output_path: str, _data: dict | None = None):
    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # スタイル
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Noto Sans JP'
    font.size = Pt(10.5)

    # ヘッダー（タイトルページ）
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('システム操作マニュアル\n\n')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x6c, 0xc0)

    run = header.add_run('Version 2.0\n')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # 改版履歴
    h = doc.add_heading('改版履歴', level=1)
    h.runs[0].font.color.rgb = RGBColor(0x00, 0x6c, 0xc0)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    headers = ['版数', '日付', '改訂内容', '担当者']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '006cc0')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        cell.paragraphs[0].runs[0].font.bold = True

    revisions = [
        ['1.0', '2025/10/01', '初版作成', '山田'],
        ['1.1', '2025/11/15', '機能追加に伴う更新', '鈴木'],
        ['2.0', '2026/01/30', 'UI変更対応', '田中'],
    ]

    for i, row_data in enumerate(revisions):
        for j, text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = text

    doc.add_paragraph()

    # 目次
    toc = doc.add_heading('目次', level=1)
    toc.runs[0].font.color.rgb = RGBColor(0x00, 0x6c, 0xc0)

    toc_items = [
        '1. はじめに',
        '2. システム概要',
        '3. インストール',
        '4. 基本操作',
        '5. トラブルシューティング',
        '6. 付録'
    ]

    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # 本文
    sections = [
        ('1. はじめに', [
            '本マニュアルでは、システムのインストールから基本操作までを詳しく説明します。',
            '対象読者：システム管理者、一般ユーザー',
            '前提条件：Windows 10以上、メモリ8GB以上'
        ]),
        ('2. システム概要', [
            '本システムは、業務効率化を目的とした統合管理プラットフォームです。',
            '主な機能：ユーザー管理、データ分析、レポート出力'
        ]),
        ('3. インストール', [
            'インストーラーをダウンロードし、以下の手順でインストールしてください。'
        ])
    ]

    for heading, contents in sections:
        h = doc.add_heading(heading, level=1)
        h.runs[0].font.color.rgb = RGBColor(0x00, 0x6c, 0xc0)

        for content in contents:
            p = doc.add_paragraph(content)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)

    # インストール手順
    h = doc.add_heading('インストール手順', level=2)
    h.runs[0].font.color.rgb = RGBColor(0x00, 0x6c, 0xc0)

    steps = [
        'インストーラーをダブルクリックして起動',
        'ライセンス契約に同意',
        'インストール先フォルダを指定',
        '「インストール」ボタンをクリック',
        '完了メッセージが表示されたら再起動'
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.paragraph_format.line_spacing = 1.5

    # コード例
    h = doc.add_heading('設定ファイル例', level=2)
    h.runs[0].font.color.rgb = RGBColor(0x00, 0x6c, 0xc0)

    doc.add_paragraph('config.iniの設定例：')

    code = '''[database]
host = localhost
port = 5432
name = production_db

[api]
endpoint = https://api.example.com
version = v2
timeout = 30'''

    add_code_block(doc, code)

    # 注意事項
    doc.add_paragraph()
    note = doc.add_paragraph()
    run = note.add_run('【注意】')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
    note.add_run('設定変更後は必ずサービスを再起動してください。')

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Created: {output}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate manual DOCX (python-docx)")
    parser.add_argument("--output", default=str(Path("output/word/02_manual.docx")))
    parser.add_argument("--data", help="Inline JSON string (optional)")
    parser.add_argument("--data-file", help="JSON file path (optional)")
    args = parser.parse_args()

    data = _load_data(args)
    create_manual(args.output, data)
