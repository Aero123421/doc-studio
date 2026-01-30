#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
プロポーザル/企画書 Wordテンプレート
- フォーマルなビジネス文書
- カラー: ネイビー + グレー
"""

import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _load_data(args: argparse.Namespace) -> dict:
    if args.data and args.data_file:
        raise SystemExit("Use either --data or --data-file (not both)")

    if args.data_file:
        p = Path(args.data_file)
        return json.loads(p.read_text(encoding="utf-8"))

    if args.data:
        return json.loads(args.data)

    return {}

def set_cell_shading(cell, color):
    """セルの背景色を設定"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_proposal(output_path: str, _data: dict | None = None):
    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # スタイル設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Noto Sans JP'
    font.size = Pt(11)

    # カバーページ
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_para.add_run('\n\n\n\n')
    run = cover_para.add_run('事業提案書\n')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    run = cover_para.add_run('\n')
    run = cover_para.add_run('新規事業立ち上げに関するご提案\n\n\n\n')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    run = cover_para.add_run('株式会社サンプル\n')
    run.font.size = Pt(12)
    run = cover_para.add_run('2026年1月')
    run.font.size = Pt(12)

    # 改ページ
    doc.add_page_break()

    # 目次
    toc_heading = doc.add_heading('目次', level=1)
    toc_heading.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    toc_items = [
        ('1. 提案の背景', '3'),
        ('2. 市場分析', '4'),
        ('3. 提案内容', '5'),
        ('4. 実施スケジュール', '6'),
        ('5. 予算案', '7'),
        ('6. 期待される効果', '8'),
    ]

    for title, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{title}').font.size = Pt(11)
        # ドットリーダー
        p.add_run('\t' + '.' * 40 + '\t').font.size = Pt(11)
        p.add_run(page).font.size = Pt(11)

    doc.add_page_break()

    # 本文
    sections = [
        ('1. 提案の背景', [
            '貴社の事業拡大戦略を拝見し、新規事業分野における当社のソリューションが貴社の成長に貢献できると確信いたしました。',
            '本事業提案書では、市場トレンドの分析に基づいた新規事業の方向性と、具体的な実施計画をご提案いたします。'
        ]),
        ('2. 市場分析', [
            '対象市場は年率15%の成長を見込む成長市場であり、早急な参入が求められています。',
            '競合分析の結果、差別化のポイントとして「品質」と「スピード」が重要であることが判明しました。'
        ]),
        ('3. 提案内容', [
            '事業コンセプト：「顧客価値を最優先にした革新的なサービス提供」',
            'ターゲット：30-40代のデジタルネイティブ層',
            '差別化戦略：AI技術を活用したパーソナライゼーション'
        ]),
    ]

    for heading, contents in sections:
        h = doc.add_heading(heading, level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

        for content in contents:
            p = doc.add_paragraph(content, style='List Bullet')
            p.paragraph_format.line_spacing = 1.5

    # 予算表
    budget_heading = doc.add_heading('5. 予算案', level=1)
    budget_heading.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    # ヘッダー
    headers = ['項目', '詳細', '金額', '備考']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '1e3a5f')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        cell.paragraphs[0].runs[0].font.bold = True

    # データ
    data = [
        ['人件費', 'エンジニア3名', '¥15,000,000', '6ヶ月分'],
        ['システム費', 'クラウド基盤', '¥3,000,000', '年間'],
        ['マーケティング', 'プロモーション', '¥5,000,000', '初期投資'],
        ['運営費', 'オフィス・備品', '¥2,000,000', '年間'],
        ['合計', '', '¥25,000,000', ''],
    ]

    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = text
            if i == len(data) - 1:  # 合計行
                set_cell_shading(cell, 'f5f5f5')
                cell.paragraphs[0].runs[0].font.bold = True

    # フッター用の段落
    doc.add_paragraph()
    doc.add_paragraph()

    # 連絡先
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('\n\nご不明点等ございましたら、下記までお問い合わせください。\n')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    run = footer.add_run('株式会社サンプル 営業部\n')
    run.font.size = Pt(10)
    run = footer.add_run('TEL: 03-1234-5678 / Email: sales@example.com')
    run.font.size = Pt(10)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Created: {output}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate proposal DOCX (python-docx)")
    parser.add_argument("--output", default=str(Path("output/word/01_proposal.docx")))
    parser.add_argument("--data", help="Inline JSON string (optional)")
    parser.add_argument("--data-file", help="JSON file path (optional)")
    args = parser.parse_args()

    data = _load_data(args)
    create_proposal(args.output, data)
