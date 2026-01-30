#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
履歴書 Wordテンプレート
- クリーンで読みやすいデザイン
- カラー: ブラック + アクセント
"""

import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _load_data(args: argparse.Namespace) -> dict:
    if args.data and args.data_file:
        raise SystemExit("Use either --data or --data-file (not both)")

    if args.data_file:
        p = Path(args.data_file)
        return json.loads(p.read_text(encoding="utf-8"))

    if args.data:
        return json.loads(args.data)

    return {}

def set_cell_shading(cell, color):
    """セルの背景色を設定"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_resume(output_path: str, _data: dict | None = None):
    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # スタイル
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Noto Sans JP'
    font.size = Pt(10.5)

    # ヘッダー（名前）
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('山田 太郎')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # 連絡先
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('〒100-0001 東京都千代田区1-1-1\n')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run = contact.add_run('TEL: 090-1234-5678 | Email: yamada@example.com')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # スペース

    # セクション関数
    def add_section(title):
        h = doc.add_heading(title, level=2)
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x5f, 0x7a)
        h.runs[0].font.size = Pt(14)
        # 下線
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('_' * 50)
        run.font.color.rgb = RGBColor(0x1a, 0x5f, 0x7a)

    # 職務経歴
    add_section('職務経歴')

    experiences = [
        {
            'company': '株式会社テックイノベーション',
            'period': '2020年4月 〜 現在',
            'position': 'シニアデベロッパー',
            'description': 'クラウドネイティブアプリケーションの設計・開発。マイクロサービスアーキテクチャの導入とチームリード。'
        },
        {
            'company': '株式会社デジタルソリューション',
            'period': '2017年4月 〜 2020年3月',
            'position': 'ソフトウェアエンジニア',
            'description': 'Webアプリケーション開発。フロントエンド・バックエンドのフルスタック開発を担当。'
        },
        {
            'company': '株式会社スタートアップ',
            'period': '2015年4月 〜 2017年3月',
            'position': 'ジュニアデベロッパー',
            'description': '新規事業の立ち上げ。モバイルアプリ開発チームに所属。'
        }
    ]

    for exp in experiences:
        # 会社名
        p = doc.add_paragraph()
        run = p.add_run(exp['company'])
        run.font.bold = True
        run.font.size = Pt(11)

        # 期間
        p = doc.add_paragraph()
        run = p.add_run(exp['period'])
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 役職
        p = doc.add_paragraph()
        run = p.add_run(exp['position'])
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1a, 0x5f, 0x7a)

        # 説明
        p = doc.add_paragraph(exp['description'])
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)

    # スキル
    add_section('スキル')

    skills = [
        ('プログラミング言語', 'Python, JavaScript, TypeScript, Go, Java'),
        ('フレームワーク', 'React, Vue.js, Django, FastAPI, Node.js'),
        ('クラウド', 'AWS, GCP, Azure, Kubernetes, Docker'),
        ('データベース', 'PostgreSQL, MySQL, MongoDB, Redis'),
        ('ツール', 'Git, CI/CD, Terraform, Ansible')
    ]

    for category, skill_list in skills:
        p = doc.add_paragraph()
        run = p.add_run(f'{category}: ')
        run.font.bold = True
        run.font.size = Pt(10)
        p.add_run(skill_list).font.size = Pt(10)

    doc.add_paragraph()

    # 学歴
    add_section('学歴')

    edu_table = doc.add_table(rows=3, cols=2)
    edu_table.style = 'Table Grid'

    education = [
        ('2013年4月 〜 2015年3月', '東京大学大学院 情報理工学系研究科 修士課程'),
        ('2009年4月 〜 2013年3月', '東京大学 工学部 情報工学科 学士'),
    ]

    for i, (period, school) in enumerate(education):
        edu_table.rows[i].cells[0].text = period
        edu_table.rows[i].cells[1].text = school
        edu_table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        edu_table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        edu_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    # 資格
    add_section('資格')

    certifications = [
        'AWS Certified Solutions Architect - Professional (2024)',
        'Google Cloud Professional Cloud Architect (2023)',
        '基本情報技術者 (2015)',
        'TOEIC 950点 (2014)'
    ]

    for cert in certifications:
        p = doc.add_paragraph(cert, style='List Bullet')
        p.paragraph_format.line_spacing = 1.5

    # 自己PR
    add_section('自己PR')

    pr_text = """技術力とコミュニケーション能力を兼ね備えたエンジニアです。
複数のプロジェクトでチームリードを経験し、メンバーのマネジメントと技術的な指導を行ってきました。
新しい技術を学ぶことに情熱を持ち、継続的な自己研鑽を大切にしています。"""

    p = doc.add_paragraph(pr_text)
    p.paragraph_format.line_spacing = 1.8

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Created: {output}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate resume DOCX (python-docx)")
    parser.add_argument("--output", default=str(Path("output/word/03_resume.docx")))
    parser.add_argument("--data", help="Inline JSON string (optional)")
    parser.add_argument("--data-file", help="JSON file path (optional)")
    args = parser.parse_args()

    data = _load_data(args)
    create_resume(args.output, data)
